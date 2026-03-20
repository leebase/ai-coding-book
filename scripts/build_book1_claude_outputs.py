#!/usr/bin/env python3
"""Build EPUB, DOCX, and PDF for the Book 1 Claude Code edition."""

from __future__ import annotations

import re
import uuid
import zipfile
from pathlib import Path

import markdown as md_lib
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import inch
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Image, PageBreak, Paragraph, Preformatted, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parent.parent
BOOK_DIR = ROOT / "book1-claude-code"
MD_PATH = BOOK_DIR / "software-engineering-with-ai-claude-code.md"
COVER_PATH = BOOK_DIR / "chapters" / "cover-optimized.png"
EPUB_PATH = BOOK_DIR / "software-engineering-with-ai-claude-code.epub"
DOCX_PATH = BOOK_DIR / "software-engineering-with-ai-claude-code.docx"
PDF_PATH = BOOK_DIR / "software-engineering-with-ai-claude-code.pdf"

TITLE = "Software Engineering with AI"
SUBTITLE = "Claude Code Edition"
AUTHOR = "Lee Harrington"
BOOK_ID = str(uuid.uuid4())

DARK = RGBColor(0x1A, 0x1A, 0x2E)
MID = RGBColor(0x2E, 0x40, 0x57)
GREY = RGBColor(0x55, 0x55, 0x55)

PAGE_SIZE = (6 * inch, 9 * inch)
MARGIN = 0.7 * inch

EPUB_CSS = """body{font-family:Georgia,"Times New Roman",serif;font-size:1em;line-height:1.65;color:#1a1a1a;margin:0 5%;}
h1{font-family:"Helvetica Neue",Arial,sans-serif;font-size:1.7em;font-weight:bold;color:#1A1A2E;margin-top:2em;margin-bottom:0.5em;page-break-before:always;}
h2{font-family:"Helvetica Neue",Arial,sans-serif;font-size:1.2em;font-weight:bold;color:#2E4057;margin-top:1.6em;margin-bottom:0.4em;}
h3{font-family:"Helvetica Neue",Arial,sans-serif;font-size:1.05em;font-weight:bold;color:#444;margin-top:1.2em;margin-bottom:0.3em;}
p{margin:0 0 0.8em 0;}
pre{background:#f5f5f5;border:1px solid #ddd;border-left:3px solid #2E4057;padding:.75em 1em;font-family:"Courier New",Courier,monospace;font-size:.85em;white-space:pre-wrap;word-wrap:break-word;margin:.8em 0;}
code{font-family:"Courier New",Courier,monospace;font-size:.88em;background:#f5f5f5;padding:.1em .3em;}
blockquote{border-left:4px solid #2E4057;background:#f0f4f8;margin:1em 0;padding:.6em 1em;color:#333;}
blockquote p{margin:.2em 0;}
ul,ol{padding-left:1.5em;margin:.5em 0;}
li{margin-bottom:.3em;}
hr{border:none;border-top:1px solid #ccc;margin:1.5em 0;}
table{border-collapse:collapse;width:100%;margin:1em 0;font-size:.9em;}
th{background:#2E4057;color:white;padding:.4em .6em;text-align:left;}
td{border:1px solid #ccc;padding:.4em .6em;}
tr:nth-child(even){background:#f9f9f9;}
"""


def split_chapters(text: str) -> list[str]:
    parts = []
    current = []
    in_code = False
    started = False

    for line in text.splitlines():
        if line.startswith("```"):
            in_code = not in_code

        is_book_heading = (
            not in_code
            and line.startswith("# ")
            and (
                line == "# Introduction"
                or line == "# Conclusion"
                or re.match(r"^# Chapter \d+:", line)
            )
        )

        if not started:
            if is_book_heading:
                started = True
                current = [line]
            continue

        if is_book_heading and current:
            parts.append("\n".join(current).strip())
            current = [line]
            continue

        current.append(line)

    if current:
        parts.append("\n".join(current).strip())

    return [p for p in parts if p]


def chapter_title(md_text: str, fallback: str) -> str:
    for line in md_text.splitlines():
        if line.startswith("# "):
            return line[2:].strip()
    return fallback


def inline_html(text: str) -> str:
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(r"`([^`]+)`", r"<font name='Courier'>\1</font>", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\*([^*]+)\*", r"<i>\1</i>", text)
    return text


def add_inline_runs(paragraph, text: str) -> None:
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def build_epub(text: str) -> None:
    chapters_md = split_chapters(text)
    chapter_titles = [chapter_title(ch, f"Chapter {i+1}") for i, ch in enumerate(chapters_md)]
    cover_bytes = COVER_PATH.read_bytes()

    manifest_items = "\n    ".join(
        f'<item id="chap{i+1:02d}" href="Text/chap_{i+1:02d}.xhtml" media-type="application/xhtml+xml"/>'
        for i in range(len(chapters_md))
    )
    spine_items = "\n    ".join(f'<itemref idref="chap{i+1:02d}"/>' for i in range(len(chapters_md)))
    opf = f"""<?xml version="1.0" encoding="utf-8"?>
<package version="3.0" xmlns="http://www.idpf.org/2007/opf" unique-identifier="bookid">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
    <dc:identifier id="bookid">urn:uuid:{BOOK_ID}</dc:identifier>
    <dc:title>{TITLE}: {SUBTITLE}</dc:title>
    <dc:creator>{AUTHOR}</dc:creator>
    <dc:language>en</dc:language>
    <dc:publisher>Leebase Press</dc:publisher>
    <dc:rights>Copyright © 2025 Lee Harrington. All rights reserved.</dc:rights>
    <meta name="cover" content="cover-image"/>
    <meta property="dcterms:modified">2025-01-01T00:00:00Z</meta>
  </metadata>
  <manifest>
    <item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>
    <item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>
    <item id="css" href="Styles/main.css" media-type="text/css"/>
    <item id="cover-image" href="Images/cover.png" media-type="image/png" properties="cover-image"/>
    <item id="cover-page" href="Text/cover.xhtml" media-type="application/xhtml+xml"/>
    {manifest_items}
  </manifest>
  <spine toc="ncx">
    <itemref idref="cover-page" linear="yes"/>
    {spine_items}
  </spine>
</package>"""

    nav_points = "\n    ".join(
        f"""<navPoint id="nav{i+1}" playOrder="{i+1}">
      <navLabel><text>{title}</text></navLabel>
      <content src="Text/chap_{i+1:02d}.xhtml"/>
    </navPoint>"""
        for i, title in enumerate(chapter_titles)
    )
    ncx = f"""<?xml version="1.0" encoding="utf-8"?>
<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">
  <head>
    <meta name="dtb:uid" content="urn:uuid:{BOOK_ID}"/>
    <meta name="dtb:depth" content="1"/>
    <meta name="dtb:totalPageCount" content="0"/>
    <meta name="dtb:maxPageNumber" content="0"/>
  </head>
  <docTitle><text>{TITLE}: {SUBTITLE}</text></docTitle>
  <navMap>
    {nav_points}
  </navMap>
</ncx>"""

    nav_links = "\n      ".join(
        f'<li><a href="Text/chap_{i+1:02d}.xhtml">{title}</a></li>'
        for i, title in enumerate(chapter_titles)
    )
    nav_xhtml = f"""<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" xml:lang="en">
<head><meta charset="UTF-8"/><title>Table of Contents</title></head>
<body>
  <nav epub:type="toc">
    <h1>Table of Contents</h1>
    <ol>
      {nav_links}
    </ol>
  </nav>
</body>
</html>"""

    cover_xhtml = """<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="en">
<head><meta charset="UTF-8"/><title>Cover</title>
<style>body{margin:0;padding:0;text-align:center;}img{max-width:100%;height:auto;}</style>
</head>
<body><img src="../Images/cover.png" alt="Cover"/></body>
</html>"""

    with zipfile.ZipFile(str(EPUB_PATH), "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("mimetype", "application/epub+zip", compress_type=zipfile.ZIP_STORED)
        zf.writestr(
            "META-INF/container.xml",
            """<?xml version="1.0"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles>
    <rootfile full-path="EPUB/content.opf" media-type="application/oebps-package+xml"/>
  </rootfiles>
</container>""",
        )
        zf.writestr("EPUB/content.opf", opf)
        zf.writestr("EPUB/toc.ncx", ncx)
        zf.writestr("EPUB/nav.xhtml", nav_xhtml)
        zf.writestr("EPUB/Styles/main.css", EPUB_CSS)
        zf.writestr("EPUB/Images/cover.png", cover_bytes)
        zf.writestr("EPUB/Text/cover.xhtml", cover_xhtml)
        for i, ch_md in enumerate(chapters_md):
            body = md_lib.markdown(ch_md, extensions=["fenced_code", "tables"])
            html = f"""<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xml:lang="en">
<head><meta charset="UTF-8"/><title>{chapter_titles[i]}</title>
<link rel="stylesheet" type="text/css" href="../Styles/main.css"/>
</head>
<body>
{body}
</body>
</html>"""
            zf.writestr(f"EPUB/Text/chap_{i+1:02d}.xhtml", html)


def build_docx(text: str) -> None:
    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(6)
        section.page_height = Inches(9)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(11)
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Arial"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = DARK
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Arial"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = MID
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Arial"
    h3.font.size = Pt(12)
    h3.font.bold = True
    if "CodeBlock" not in [style.name for style in doc.styles]:
        style = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Courier New"
        style.font.size = Pt(9)
        style.paragraph_format.left_indent = Inches(0.3)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(4)

    core = doc.core_properties
    core.title = f"{TITLE}: {SUBTITLE}"
    core.author = AUTHOR
    core.subject = "AI-assisted software engineering guide"
    core.comments = "Generated from software-engineering-with-ai-claude-code.md"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(COVER_PATH), width=Inches(5.4))
    doc.add_section(WD_SECTION.NEW_PAGE)
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = DARK
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(SUBTITLE)
    run.font.name = "Arial"
    run.font.size = Pt(13)
    run.italic = True
    run.font.color.rgb = GREY
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_before = Pt(20)
    run = author.add_run(AUTHOR)
    run.font.name = "Arial"
    run.font.size = Pt(15)
    run.font.color.rgb = MID
    doc.add_section(WD_SECTION.NEW_PAGE)

    lines = text.splitlines()
    in_code = False
    code_lines = []
    in_callout = False
    callout_lines = []

    def flush_code() -> None:
        nonlocal code_lines
        if not code_lines:
            return
        for line in code_lines:
            doc.add_paragraph(line if line else " ", style="CodeBlock")
        code_lines = []

    def flush_callout() -> None:
        nonlocal callout_lines
        if not callout_lines:
            return
        for raw in callout_lines:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.3)
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(3)
            add_inline_runs(paragraph, re.sub(r"^>\s*", "", raw))
        callout_lines = []

    for line in lines:
        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_callout()
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.startswith(">"):
            in_callout = True
            callout_lines.append(line)
            continue
        if in_callout:
            flush_callout()
            in_callout = False
        stripped = line.strip()
        if not stripped:
            continue
        if stripped == "---":
            doc.add_paragraph()
            continue
        if stripped.startswith("# "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline_runs(p, stripped[2:])
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline_runs(p, stripped[3:])
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline_runs(p, stripped[4:])
            continue
        if re.match(r"^\d+\. ", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, re.sub(r"^\d+\. ", "", stripped))
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:])
            continue
        if stripped.startswith("|"):
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        add_inline_runs(p, stripped)

    if in_callout:
        flush_callout()
    if in_code:
        flush_code()

    doc.save(str(DOCX_PATH))


def build_pdf(text: str) -> None:
    styles = getSampleStyleSheet()
    styles["BodyText"].fontName = "Times-Roman"
    styles["BodyText"].fontSize = 11
    styles["BodyText"].leading = 15
    styles["Heading1"].fontName = "Helvetica-Bold"
    styles["Heading1"].fontSize = 18
    styles["Heading1"].leading = 22
    styles["Heading1"].spaceBefore = 18
    styles["Heading1"].spaceAfter = 10
    styles["Heading2"].fontName = "Helvetica-Bold"
    styles["Heading2"].fontSize = 14
    styles["Heading2"].leading = 18
    styles["Heading2"].spaceBefore = 14
    styles["Heading2"].spaceAfter = 8
    styles["Heading3"].fontName = "Helvetica-Bold"
    styles["Heading3"].fontSize = 12
    styles["Heading3"].leading = 15
    styles["Heading3"].spaceBefore = 10
    styles["Heading3"].spaceAfter = 6
    styles.add(ParagraphStyle(name="TitlePage", parent=styles["Heading1"], alignment=TA_CENTER, fontSize=24, leading=28))
    styles.add(ParagraphStyle(name="Subtitle", parent=styles["BodyText"], alignment=TA_CENTER, fontName="Helvetica-Oblique", fontSize=13, textColor=colors.HexColor("#555555")))
    styles.add(ParagraphStyle(name="Author", parent=styles["BodyText"], alignment=TA_CENTER, fontName="Helvetica", fontSize=15, textColor=colors.HexColor("#2E4057")))
    styles.add(ParagraphStyle(name="Callout", parent=styles["BodyText"], leftIndent=18, borderPadding=6, borderColor=colors.HexColor("#2E4057"), borderWidth=1, borderLeft=True, backColor=colors.HexColor("#F0F4F8"), spaceBefore=4, spaceAfter=4))

    story = [
        Image(str(COVER_PATH), width=5.2 * inch, height=6.93 * inch),
        PageBreak(),
        Spacer(1, 1.5 * inch),
        Paragraph(TITLE, styles["TitlePage"]),
        Spacer(1, 0.2 * inch),
        Paragraph(SUBTITLE, styles["Subtitle"]),
        Spacer(1, 0.4 * inch),
        Paragraph(AUTHOR, styles["Author"]),
        PageBreak(),
    ]

    lines = text.splitlines()
    in_code = False
    code_lines = []
    callout_lines = []

    def flush_code():
        nonlocal code_lines
        if code_lines:
            story.append(
                Preformatted(
                    "\n".join(code_lines),
                    ParagraphStyle(
                        "CodeBlock",
                        fontName="Courier",
                        fontSize=8.5,
                        leading=10,
                        leftIndent=16,
                        rightIndent=16,
                        backColor=colors.HexColor("#F5F5F5"),
                        borderColor=colors.HexColor("#DDDDDD"),
                        borderWidth=1,
                        borderPadding=6,
                    ),
                )
            )
            story.append(Spacer(1, 0.08 * inch))
            code_lines = []

    def flush_callout():
        nonlocal callout_lines
        if callout_lines:
            for raw in callout_lines:
                story.append(Paragraph(inline_html(re.sub(r"^>\s*", "", raw)), styles["Callout"]))
            story.append(Spacer(1, 0.08 * inch))
            callout_lines = []

    for line in lines:
        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_callout()
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.startswith(">"):
            callout_lines.append(line)
            continue
        if callout_lines:
            flush_callout()
        stripped = line.strip()
        if not stripped:
            continue
        if stripped == "---":
            story.append(Spacer(1, 0.12 * inch))
            continue
        if stripped.startswith("# "):
            story.append(Paragraph(inline_html(stripped[2:]), styles["Heading1"]))
            continue
        if stripped.startswith("## "):
            story.append(Paragraph(inline_html(stripped[3:]), styles["Heading2"]))
            continue
        if stripped.startswith("### "):
            story.append(Paragraph(inline_html(stripped[4:]), styles["Heading3"]))
            continue
        if re.match(r"^\d+\. ", stripped):
            story.append(Paragraph(f"&#8226; {inline_html(re.sub(r'^\d+\. ', '', stripped))}", styles["BodyText"]))
            continue
        if stripped.startswith("- "):
            story.append(Paragraph(f"&#8226; {inline_html(stripped[2:])}", styles["BodyText"]))
            continue
        if stripped.startswith("|"):
            continue
        story.append(Paragraph(inline_html(stripped), styles["BodyText"]))

    if callout_lines:
        flush_callout()
    if code_lines:
        flush_code()

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=PAGE_SIZE,
        leftMargin=MARGIN,
        rightMargin=MARGIN,
        topMargin=MARGIN,
        bottomMargin=MARGIN,
        title=f"{TITLE}: {SUBTITLE}",
        author=AUTHOR,
    )
    doc.build(story)


def main() -> None:
    text = MD_PATH.read_text(encoding="utf-8")
    build_epub(text)
    build_docx(text)
    build_pdf(text)
    for path in (EPUB_PATH, DOCX_PATH, PDF_PATH):
        print(f"Written: {path} ({path.stat().st_size // 1024}KB)")


if __name__ == "__main__":
    main()
