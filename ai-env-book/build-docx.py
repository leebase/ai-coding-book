#!/usr/bin/env python3
"""Build DOCX for Your Dev Environment: A Guide for AI-Assisted Developers."""

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

BOOK_DIR = Path(__file__).parent
MD_PATH = BOOK_DIR / "your-dev-environment.md"
COVER_PATH = BOOK_DIR / "DevEnvCover.png"
OUT_PATH = BOOK_DIR / "your-dev-environment.docx"

TITLE = "Your Dev Environment"
SUBTITLE = "A Guide for AI-Assisted Developers"
AUTHOR = "Lee Harrington"

DARK = RGBColor(0x1A, 0x1A, 0x2E)
MID = RGBColor(0x2E, 0x40, 0x57)
GREY = RGBColor(0x55, 0x55, 0x55)


def set_page_layout(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(6)
        section.page_height = Inches(9)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)


def ensure_code_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(11)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Arial"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = DARK

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Arial"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = MID

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Arial"
    h3.font.size = Pt(12)
    h3.font.bold = True

    if "CodeBlock" not in [style.name for style in doc.styles]:
        style = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Courier New"
        style.font.size = Pt(9)
        style.paragraph_format.left_indent = Inches(0.3)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(4)


def add_title_pages(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(COVER_PATH), width=Inches(5.4))

    doc.add_section(WD_SECTION.NEW_PAGE)

    for _ in range(3):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = DARK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(SUBTITLE)
    run.font.name = "Arial"
    run.font.size = Pt(13)
    run.italic = True
    run.font.color.rgb = GREY

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_before = Pt(20)
    run = author.add_run(AUTHOR)
    run.font.name = "Arial"
    run.font.size = Pt(15)
    run.font.color.rgb = MID

    doc.add_section(WD_SECTION.NEW_PAGE)


def add_inline_runs(paragraph, text: str) -> None:
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def render_markdown(doc: Document, md_text: str) -> None:
    lines = md_text.splitlines()
    in_code = False
    code_lines = []
    in_callout = False
    callout_lines = []

    def flush_code() -> None:
        nonlocal code_lines
        if not code_lines:
            return
        for line in code_lines:
            doc.add_paragraph(line if line else " ", style="CodeBlock")
        code_lines = []

    def flush_callout() -> None:
        nonlocal callout_lines
        if not callout_lines:
            return
        for raw in callout_lines:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.3)
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(3)
            add_inline_runs(paragraph, re.sub(r"^>\s*", "", raw))
        callout_lines = []

    for line in lines:
        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_callout()
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.startswith(">"):
            in_callout = True
            callout_lines.append(line)
            continue

        if in_callout:
            flush_callout()
            in_callout = False

        stripped = line.strip()
        if not stripped:
            continue

        if stripped == "---":
            doc.add_paragraph()
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline_runs(p, stripped[2:])
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline_runs(p, stripped[3:])
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline_runs(p, stripped[4:])
            continue

        if re.match(r"^\d+\. ", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, re.sub(r"^\d+\. ", "", stripped))
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:])
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        add_inline_runs(p, stripped)

    if in_callout:
        flush_callout()
    if in_code:
        flush_code()


def main() -> None:
    doc = Document()
    set_page_layout(doc)
    ensure_code_style(doc)
    core = doc.core_properties
    core.title = f"{TITLE}: {SUBTITLE}"
    core.author = AUTHOR
    core.subject = "AI-assisted development environment guide"
    core.comments = "Generated from your-dev-environment.md"
    add_title_pages(doc)
    render_markdown(doc, MD_PATH.read_text(encoding="utf-8"))
    doc.save(str(OUT_PATH))
    size = OUT_PATH.stat().st_size
    print(f"Written: {OUT_PATH} ({size // 1024}KB)")


if __name__ == "__main__":
    main()
